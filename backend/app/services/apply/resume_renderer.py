"""
Server-side, ATS-safe resume + cover-letter renderer.

Consumes the structured `resume` spec produced by
`resume_tailor_llm.tailor_resume` (name / contact / summary / skills /
experience / education / certifications / projects) and renders it to a DOCX
and a PDF that pass ATS parsing: a single column, standard fonts, real headings,
no tables, no text boxes, no images, no multi-column layout. This is the
production implementation behind the previously-stubbed
`render_and_store_tailored` hook, so every application gets a freshly tailored
resume + cover letter document.

Original PlaceUp templates — no third-party code or branding is used.

Renderers are pure `spec -> bytes` functions with no storage/network, so they
are unit-testable and safe to run inside a Cloud Run request or job.
"""
from __future__ import annotations

import io
import logging
from typing import Any, Optional

log = logging.getLogger("placeup.apply.render")

# --- shared style constants (ATS-safe, single column) ---
_FONT = "Calibri"
_NAME_PT = 20
_SECTION_PT = 12
_BODY_PT = 10.5


def _as_list(value: Any) -> list:
    if isinstance(value, list):
        return value
    if value in (None, ""):
        return []
    return [value]


def _contact_line(resume: dict) -> str:
    parts = [str(p).strip() for p in _as_list(resume.get("contact")) if str(p).strip()]
    return "  |  ".join(parts)


# ─────────────────────────── DOCX ───────────────────────────

def render_resume_docx(resume: dict) -> bytes:
    """Render the tailored resume spec to an ATS-safe .docx (bytes)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(_BODY_PT)

    # Tighter margins so a tailored one-pager fits.
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(48)

    def _heading(text: str) -> None:
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(_SECTION_PT)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
        # a thin rule under the heading via bottom border
        _bottom_border(p)

    def _bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(str(text))

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nrun = name_p.add_run(str(resume.get("name") or "").strip())
    nrun.bold = True
    nrun.font.size = Pt(_NAME_PT)

    # Contact
    contact = _contact_line(resume)
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(contact).font.size = Pt(_BODY_PT)

    # Summary
    summary = str(resume.get("summary") or "").strip()
    if summary:
        _heading("Summary")
        doc.add_paragraph(summary)

    # Core Skills
    skills = _as_list(resume.get("skills"))
    if skills:
        _heading("Core Skills")
        for group in skills:
            if isinstance(group, dict):
                cat = str(group.get("category") or "").strip()
                items = ", ".join(str(i).strip() for i in _as_list(group.get("items")) if str(i).strip())
                line = f"{cat}: {items}" if cat else items
            else:
                line = str(group)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if isinstance(group, dict) and group.get("category"):
                p.add_run(f"{group['category']}: ").bold = True
                p.add_run(items)
            else:
                p.add_run(line)

    # Professional Experience
    experience = _as_list(resume.get("experience"))
    if experience:
        _heading("Professional Experience")
        for role in experience:
            if not isinstance(role, dict):
                continue
            head = doc.add_paragraph()
            head.paragraph_format.space_after = Pt(0)
            title = str(role.get("title") or "").strip()
            company = str(role.get("company") or "").strip()
            left = title + (f", {company}" if company else "")
            run = head.add_run(left)
            run.bold = True
            dates = str(role.get("dates") or "").strip()
            loc = str(role.get("location") or "").strip()
            right = "  ".join(x for x in (loc, dates) if x)
            if right:
                head.add_run(f"   ({right})").italic = True
            for b in _as_list(role.get("bullets")):
                if str(b).strip():
                    _bullet(b)

    # Education
    education = _as_list(resume.get("education"))
    if education:
        _heading("Education")
        for ed in education:
            if not isinstance(ed, dict):
                doc.add_paragraph(str(ed))
                continue
            p = doc.add_paragraph()
            deg = str(ed.get("degree") or "").strip()
            inst = str(ed.get("institution") or "").strip()
            p.add_run(deg).bold = True
            tail = "  ".join(x for x in (inst, str(ed.get("location") or "").strip(), str(ed.get("dates") or "").strip()) if x)
            if tail:
                p.add_run(f" — {tail}")

    # Certifications
    certs = [str(c).strip() for c in _as_list(resume.get("certifications")) if str(c).strip()]
    if certs:
        _heading("Certifications")
        for c in certs:
            _bullet(c)

    # Projects
    projects = [str(pj).strip() for pj in _as_list(resume.get("projects")) if str(pj).strip()]
    if projects:
        _heading("Projects")
        for pj in projects:
            _bullet(pj)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _bottom_border(paragraph) -> None:
    """Add a thin bottom border to a paragraph (section-rule) without tables."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9AA6B2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ─────────────────────────── PDF ───────────────────────────

def render_resume_pdf(resume: dict) -> bytes:
    """Render the tailored resume spec to an ATS-safe single-column PDF."""
    return _render_pdf(resume, kind="resume")


def _styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib import colors

    ss = getSampleStyleSheet()
    styles = {
        "name": ParagraphStyle("puName", parent=ss["Title"], fontSize=_NAME_PT,
                               alignment=TA_CENTER, spaceAfter=2, textColor=colors.HexColor("#111827")),
        "contact": ParagraphStyle("puContact", parent=ss["Normal"], fontSize=_BODY_PT,
                                  alignment=TA_CENTER, textColor=colors.HexColor("#374151"), spaceAfter=6),
        "section": ParagraphStyle("puSection", parent=ss["Heading2"], fontSize=_SECTION_PT,
                                  textColor=colors.HexColor("#1F2A44"), spaceBefore=8, spaceAfter=2),
        "body": ParagraphStyle("puBody", parent=ss["Normal"], fontSize=_BODY_PT, leading=13.5),
        "role": ParagraphStyle("puRole", parent=ss["Normal"], fontSize=_BODY_PT, leading=13.5, spaceBefore=4),
        "bullet": ParagraphStyle("puBullet", parent=ss["Normal"], fontSize=_BODY_PT, leading=13,
                                 leftIndent=12, bulletIndent=2, spaceAfter=1),
    }
    return styles


def _render_pdf(resume: dict, kind: str = "resume") -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from xml.sax.saxutils import escape

    styles = _styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=0.5 * inch, bottomMargin=0.5 * inch,
                            leftMargin=0.65 * inch, rightMargin=0.65 * inch, title="Resume")
    flow = []

    def section(title):
        flow.append(Paragraph(escape(title.upper()), styles["section"]))
        flow.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#9AA6B2"), spaceAfter=4))

    flow.append(Paragraph(escape(str(resume.get("name") or "")), styles["name"]))
    contact = _contact_line(resume)
    if contact:
        flow.append(Paragraph(escape(contact), styles["contact"]))

    if str(resume.get("summary") or "").strip():
        section("Summary")
        flow.append(Paragraph(escape(str(resume["summary"]).strip()), styles["body"]))

    skills = _as_list(resume.get("skills"))
    if skills:
        section("Core Skills")
        for group in skills:
            if isinstance(group, dict):
                cat = escape(str(group.get("category") or "").strip())
                items = escape(", ".join(str(i).strip() for i in _as_list(group.get("items")) if str(i).strip()))
                text = f"<b>{cat}:</b> {items}" if cat else items
            else:
                text = escape(str(group))
            flow.append(Paragraph(text, styles["body"]))

    experience = _as_list(resume.get("experience"))
    if experience:
        section("Professional Experience")
        for role in experience:
            if not isinstance(role, dict):
                continue
            title = escape(str(role.get("title") or "").strip())
            company = escape(str(role.get("company") or "").strip())
            dates = escape(str(role.get("dates") or "").strip())
            loc = escape(str(role.get("location") or "").strip())
            right = "  ".join(x for x in (loc, dates) if x)
            head = f"<b>{title}{', ' + company if company else ''}</b>"
            if right:
                head += f"   <i>({right})</i>"
            flow.append(Paragraph(head, styles["role"]))
            for b in _as_list(role.get("bullets")):
                if str(b).strip():
                    flow.append(Paragraph(escape(str(b).strip()), styles["bullet"], bulletText="•"))

    education = _as_list(resume.get("education"))
    if education:
        section("Education")
        for ed in education:
            if isinstance(ed, dict):
                deg = escape(str(ed.get("degree") or "").strip())
                tail = "  ".join(x for x in (str(ed.get("institution") or "").strip(),
                                             str(ed.get("location") or "").strip(),
                                             str(ed.get("dates") or "").strip()) if x)
                flow.append(Paragraph(f"<b>{deg}</b>{' — ' + escape(tail) if tail else ''}", styles["body"]))
            else:
                flow.append(Paragraph(escape(str(ed)), styles["body"]))

    certs = [str(c).strip() for c in _as_list(resume.get("certifications")) if str(c).strip()]
    if certs:
        section("Certifications")
        for c in certs:
            flow.append(Paragraph(escape(c), styles["bullet"], bulletText="•"))

    projects = [str(p).strip() for p in _as_list(resume.get("projects")) if str(p).strip()]
    if projects:
        section("Projects")
        for p in projects:
            flow.append(Paragraph(escape(p), styles["bullet"], bulletText="•"))

    doc.build(flow)
    return buf.getvalue()


# ─────────────────────── Cover letter ───────────────────────

def render_cover_letter_docx(body: str, *, name: str = "", contact: str = "",
                             company: str = "", role: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.name = _FONT
    doc.styles["Normal"].font.size = Pt(_BODY_PT)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(54)
        section.left_margin = section.right_margin = Pt(64)

    if name:
        h = doc.add_paragraph()
        r = h.add_run(name)
        r.bold = True
        r.font.size = Pt(14)
    if contact:
        doc.add_paragraph(contact)
    doc.add_paragraph("")
    if company or role:
        doc.add_paragraph(f"Re: {role}{' at ' + company if company else ''}".strip())
        doc.add_paragraph("")
    for para in str(body or "").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_cover_letter_pdf(body: str, *, name: str = "", contact: str = "",
                            company: str = "", role: str = "") -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from xml.sax.saxutils import escape

    styles = _styles()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, topMargin=0.75 * inch, bottomMargin=0.75 * inch,
                            leftMargin=0.9 * inch, rightMargin=0.9 * inch, title="Cover Letter")
    flow = []
    if name:
        flow.append(Paragraph(f"<b>{escape(name)}</b>", styles["role"]))
    if contact:
        flow.append(Paragraph(escape(contact), styles["contact"]))
    flow.append(Spacer(1, 10))
    if company or role:
        flow.append(Paragraph(escape(f"Re: {role}{' at ' + company if company else ''}".strip()), styles["body"]))
        flow.append(Spacer(1, 8))
    for para in str(body or "").split("\n\n"):
        if para.strip():
            flow.append(Paragraph(escape(para.strip()), styles["body"]))
            flow.append(Spacer(1, 6))
    doc.build(flow)
    return buf.getvalue()


def render_all(resume: dict, cover_letter: Optional[str] = None) -> dict:
    """Convenience: render resume (docx+pdf) and optional cover letter.
    Returns a dict of name -> bytes. Missing renderers degrade to skipping."""
    out: dict[str, bytes] = {}
    try:
        out["resume.docx"] = render_resume_docx(resume)
    except Exception as exc:  # pragma: no cover
        log.warning("resume docx render failed: %s", exc)
    try:
        out["resume.pdf"] = render_resume_pdf(resume)
    except Exception as exc:  # pragma: no cover
        log.warning("resume pdf render failed: %s", exc)
    if cover_letter:
        meta = {
            "name": str(resume.get("name") or ""),
            "contact": _contact_line(resume),
        }
        try:
            out["cover_letter.docx"] = render_cover_letter_docx(cover_letter, **meta)
        except Exception as exc:  # pragma: no cover
            log.warning("cover docx render failed: %s", exc)
        try:
            out["cover_letter.pdf"] = render_cover_letter_pdf(cover_letter, **meta)
        except Exception as exc:  # pragma: no cover
            log.warning("cover pdf render failed: %s", exc)
    return out
