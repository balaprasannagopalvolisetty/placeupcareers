"""
PlaceUp Career — Resume Parser Service
Extracts text from PDF and DOCX resume files.
"""

import io
import logging
import re
from typing import Optional
from urllib.parse import urlparse
from zipfile import ZipFile, is_zipfile

logger = logging.getLogger(__name__)

MAX_PDF_PAGES = 8
# We extract TEXT ONLY here — PyPDF2 never renders or executes the PDF — and the
# original file is NOT stored or re-served. So benign-but-common markers that
# appear in ordinary resumes exported from Word, Canva, LaTeX, Google Docs, etc.
# must NOT cause a rejection:
#   /OpenAction  -> initial view/zoom ("fit page") set on open
#   /AA          -> additional-action triggers on links/form fields
#   /EmbeddedFile-> attached fonts or the source document
# Blocking those produced false positives like "PDF contains active or embedded
# content and cannot be accepted" on legitimate resumes. We now only reject
# clear script/launch actions.
PDF_ACTIVE_CONTENT_MARKERS = (
    b"/JavaScript",
    b"/Launch",
)
DOCX_BLOCKED_PARTS = (
    "vbaProject.bin",
    "activeX/",
    "embeddings/",
)
DOCX_BLOCKED_REL_MARKERS = (
    b"oleObject",
    b"activeXControl",
)


async def parse_resume_file(
    file_content: bytes,
    filename: str,
) -> dict:
    """Parse a resume file and extract raw text + metadata.

    Supports PDF and DOCX file formats. Cleans extracted text
    and provides basic document statistics.

    Args:
        file_content: Raw file bytes
        filename: Original filename (used to detect format)

    Returns:
        Dict with keys: text, word_count, page_count, format
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if "\x00" in filename or "/" in filename or "\\" in filename:
        raise ValueError("Invalid filename.")

    if ext == "pdf":
        return await _parse_pdf(file_content)
    elif ext == "docx":
        return await _parse_docx(file_content)
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Please upload a PDF or DOCX file.")


async def _parse_pdf(content: bytes) -> dict:
    """Extract text from a PDF file using PyPDF2.

    Handles multi-page PDFs, removing headers/footers
    and normalizing whitespace.

    Args:
        content: Raw PDF bytes

    Returns:
        Dict with text, word_count, page_count
    """
    if not content.startswith(b"%PDF"):
        raise ValueError("Uploaded file is not a valid PDF.")
    lower_probe = content[: min(len(content), 2_000_000)]
    if any(marker in lower_probe for marker in PDF_ACTIVE_CONTENT_MARKERS):
        raise ValueError("PDF contains active or embedded content and cannot be accepted.")
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        if getattr(reader, "is_encrypted", False):
            raise ValueError("Encrypted PDFs are not supported.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError(f"Resume PDF is too long. Please upload {MAX_PDF_PAGES} pages or fewer.")
        pages = []
        links = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
            links.extend(_pdf_page_links(page))

        full_text = "\n\n".join(pages)
        full_text = _clean_resume_text(full_text)

        return {
            "text": full_text,
            "word_count": len(full_text.split()),
            "page_count": len(reader.pages),
            "format": "pdf",
            "links": sorted(set(links))[:20],
        }

    except ImportError:
        logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


async def _parse_docx(content: bytes) -> dict:
    """Extract text from a DOCX file using python-docx.

    Extracts text from paragraphs and tables.

    Args:
        content: Raw DOCX bytes

    Returns:
        Dict with text, word_count, page_count
    """
    if not content.startswith(b"PK") or not is_zipfile(io.BytesIO(content)):
        raise ValueError("Uploaded file is not a valid DOCX document.")
    _scan_docx_for_active_content(content)
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = []

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        links = _docx_links(content)

        full_text = "\n".join(paragraphs)
        full_text = _clean_resume_text(full_text)

        return {
            "text": full_text,
            "word_count": len(full_text.split()),
            "page_count": max(1, len(full_text) // 3000),  # Estimate
            "format": "docx",
            "links": sorted(set(links))[:20],
        }

    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def _scan_docx_for_active_content(content: bytes) -> None:
    """Reject DOCX packages with macros or embedded active objects.

    Ordinary resume hyperlinks are allowed: we parse their target strings but
    never fetch or execute them.
    """
    try:
        with ZipFile(io.BytesIO(content)) as zf:
            names = zf.namelist()
            lower_names = [name.replace("\\", "/").lower() for name in names]
            for blocked in DOCX_BLOCKED_PARTS:
                if any(blocked.lower() in name for name in lower_names):
                    raise ValueError("DOCX contains active or embedded content and cannot be accepted.")
            for name in names:
                lower = name.lower()
                if not lower.endswith(".rels") and "document.xml" not in lower:
                    continue
                data = zf.read(name)[:2_000_000]
                if any(marker in data for marker in DOCX_BLOCKED_REL_MARKERS):
                    raise ValueError("DOCX contains external or embedded content and cannot be accepted.")
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"DOCX security scan failed: {exc}")


def _pdf_page_links(page) -> list[str]:
    links: list[str] = []
    try:
        annotations = page.get("/Annots") or []
        for annot_ref in annotations:
            annot = annot_ref.get_object()
            action = annot.get("/A") or {}
            uri = action.get("/URI")
            clean = _clean_link(str(uri or ""))
            if clean:
                links.append(clean)
    except Exception:
        return links
    return links


def _docx_links(content: bytes) -> list[str]:
    links: list[str] = []
    try:
        with ZipFile(io.BytesIO(content)) as zf:
            for name in zf.namelist():
                if not name.lower().endswith(".rels"):
                    continue
                data = zf.read(name).decode("utf-8", "ignore")
                for target in re.findall(r'Target="([^"]+)"', data):
                    clean = _clean_link(target)
                    if clean:
                        links.append(clean)
    except Exception:
        return links
    return links


def _clean_link(value: str) -> str:
    text = str(value or "").strip().strip(".,;:)]}>\"'")
    if not text:
        return ""
    if text.startswith("mailto:"):
        return text
    if re.match(r"^(?:https?://|www\.|(?:linkedin|github)\.com/)", text, re.I):
        if not re.match(r"^https?://", text, re.I):
            text = "https://" + text
        parsed = urlparse(text)
        if parsed.scheme in {"http", "https"} and parsed.netloc:
            return text
    return ""


def _clean_resume_text(text: str) -> str:
    """Clean and normalize extracted resume text.

    Removes common PDF artifacts like page numbers,
    excessive whitespace, and control characters, then repairs
    "word-shattered" extractions (one word per line) that some PDF
    generators produce — those previously broke sectioning, keyword
    extraction, and the Profile resume preview.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    # Remove page numbers (e.g., "Page 1 of 3", "1/3", standalone numbers)
    text = re.sub(r"(?i)page\s+\d+\s*(of\s*\d+)?", "", text)
    text = re.sub(r"^\s*\d+\s*/\s*\d+\s*$", "", text, flags=re.MULTILINE)

    # Remove control characters
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Normalize whitespace (preserve paragraph breaks)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return _reconstruct_shattered_text(text.strip())


# Section headings recognized both as standalone lines and as ALL-CAPS tokens
# inside shattered text.
_SECTION_HEADING_TOKENS = frozenset({
    "summary", "profile", "objective",
    "skills", "competencies", "technologies", "expertise",
    "experience", "employment",
    "education", "academics",
    "projects", "portfolio",
    "certifications", "licenses", "credentials",
    "achievements", "awards", "publications",
})
_HEADING_MODIFIER_TOKENS = frozenset({
    "technical", "core", "key", "relevant", "work", "professional",
    "career", "academic", "personal", "selected",
})
_BULLET_CHARS = "•●▪‣◦∙·"


def _reconstruct_shattered_text(text: str) -> str:
    """Rejoin word-per-line PDF extractions into readable lines.

    Some PDF generators position every word separately, so extract_text()
    emits one word per line. Detection: most non-empty lines hold ≤2 short
    words. Repair: stream the words back together, starting a new line at
    section headings (ALL-CAPS tokens like EXPERIENCE / TECHNICAL SKILLS)
    and at bullet markers (kept as "• " prefixes so the UI can render
    real bullets). Normal multi-word extractions pass through untouched.
    """
    lines = [ln.strip() for ln in (text or "").splitlines()]
    nonempty = [ln for ln in lines if ln]
    if len(nonempty) < 15:
        return text
    shattered = sum(1 for ln in nonempty if len(ln.split()) <= 2 and len(ln) <= 24)
    if shattered / len(nonempty) < 0.55:
        return text

    tokens: list[str] = []
    for ln in nonempty:
        tokens.extend(ln.split())

    out: list[str] = []
    buf: list[str] = []

    def flush() -> None:
        if buf:
            out.append(" ".join(buf))
            buf.clear()

    for tok in tokens:
        stripped = tok.strip(":;,.")
        bare = stripped.lower()

        # ALL-CAPS section headings start a new line (and absorb a leading
        # modifier so "TECHNICAL SKILLS" / "WORK EXPERIENCE" stay intact).
        if bare in _SECTION_HEADING_TOKENS and stripped.isupper() and len(stripped) >= 5:
            modifier = ""
            if buf:
                tail = buf[-1].strip(":;,.")
                if tail.isupper() and tail.lower() in _HEADING_MODIFIER_TOKENS:
                    buf.pop()
                    modifier = tail + " "
            flush()
            out.append(modifier + stripped)
            continue

        # Bullet markers start a new bullet line.
        if tok and all(ch in _BULLET_CHARS for ch in tok):
            flush()
            buf.append("•")
            continue
        if tok[0] in _BULLET_CHARS:
            flush()
            buf.append("•")
            rest = tok.lstrip(_BULLET_CHARS + " ")
            if rest:
                buf.append(rest)
            continue

        buf.append(tok)
    flush()
    return "\n".join(out)


def _section_lines(text: str) -> dict[str, list[str]]:
    """Split resume text into named sections.

    Content before the first recognized heading lands in "header" (name +
    contact lines) instead of polluting the summary. Headings match with or
    without trailing colons, in any case, and with common modifier words.
    """
    headings = {
        "summary": r"(?:professional\s+|career\s+|executive\s+)?(?:summary|profile|objective)",
        "skills": r"(?:technical\s+|core\s+|key\s+|relevant\s+)?(?:skills|competencies|technologies|areas of expertise|expertise)",
        "experience": r"(?:work\s+|professional\s+|relevant\s+|career\s+)?(?:experience|employment(?:\s+history)?|work history)",
        "education": r"(?:education(?:al)?)(?:\s+background)?|academics?|academic background",
        "projects": r"(?:selected\s+|personal\s+|academic\s+|key\s+)?projects?|portfolio",
        "certifications": r"certifications?(?:\s*&?\s*licenses?)?|licenses?|credentials|certificates",
    }
    sections: dict[str, list[str]] = {key: [] for key in headings}
    sections["header"] = []
    current = "header"
    for raw in text.splitlines():
        line = raw.strip(" -\t")
        if not line:
            continue
        probe = line.strip(" :·|").strip()
        matched = None
        for key, pattern in headings.items():
            if re.fullmatch(pattern, probe, flags=re.I):
                matched = key
                break
        if matched:
            current = matched
            continue
        sections.setdefault(current, []).append(line)
    return {key: value[:40] for key, value in sections.items() if value}


def _first_match(pattern: str, text: str) -> str:
    match = re.search(pattern, text, flags=re.I)
    return match.group(0).strip() if match else ""


# ── Structured experience parsing ────────────────────────────────────────────
#
# Splits the experience section into individual positions with title, company,
# dates, and computed duration — previously all positions were rendered as one
# undifferentiated blob and Past Companies had no titles or durations.

_MONTHS = "jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec"
_DATE_TOKEN = rf"(?:(?:{_MONTHS})[a-z]*\.?\s+\d{{4}}|\d{{1,2}}/\d{{4}}|\d{{4}})"
_DATE_RANGE = re.compile(
    rf"({_DATE_TOKEN})\s*(?:[-–—~]|to|until)\s*({_DATE_TOKEN}|present|current|now|ongoing)",
    re.I,
)
_TITLE_WORDS = re.compile(
    r"\b(engineer|developer|analyst|manager|consultant|intern|architect|designer|"
    r"scientist|administrator|specialist|lead|director|officer|coordinator|"
    r"technician|associate|assistant|accountant|auditor|recruiter|representative|"
    r"support|advisor|supervisor|head|founder|owner|president|nurse|therapist)\b",
    re.I,
)
_MONTH_INDEX = {m: i + 1 for i, m in enumerate(
    ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"])}


def _parse_date_token(token: str) -> Optional[tuple[int, int]]:
    """'(month, year)' from 'Jan 2024' / '01/2024' / '2024'; None if unparseable."""
    tok = (token or "").strip().lower().rstrip(".")
    if tok in {"present", "current", "now", "ongoing"}:
        from datetime import datetime as _dt
        now = _dt.utcnow()
        return (now.month, now.year)
    m = re.match(rf"({_MONTHS})[a-z]*\.?\s+(\d{{4}})", tok)
    if m:
        return (_MONTH_INDEX.get(m.group(1)[:3], 1), int(m.group(2)))
    m = re.match(r"(\d{1,2})/(\d{4})", tok)
    if m:
        return (max(1, min(12, int(m.group(1)))), int(m.group(2)))
    m = re.match(r"(\d{4})$", tok)
    if m:
        return (1, int(m.group(1)))
    return None


def _duration_label(start: Optional[tuple[int, int]], end: Optional[tuple[int, int]]) -> str:
    if not start or not end:
        return ""
    months = (end[1] - start[1]) * 12 + (end[0] - start[0]) + 1
    if months <= 0:
        return ""
    years, rem = divmod(months, 12)
    parts = []
    if years:
        parts.append(f"{years} yr" + ("s" if years > 1 else ""))
    if rem:
        parts.append(f"{rem} mo" + ("s" if rem > 1 else ""))
    return " ".join(parts)


def _duration_label_from_dates(dates: str) -> str:
    """Human duration ("1 yr 9 mos") computed from a date-range string."""
    m = re.match(
        rf"\s*({_DATE_TOKEN})\s*(?:[-–—~]|to|until)\s*({_DATE_TOKEN}|present|current|now|ongoing)",
        str(dates or ""), re.I,
    )
    if not m:
        return ""
    return _duration_label(_parse_date_token(m.group(1)), _parse_date_token(m.group(2)))


_MONTH = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
_DATE_RANGE_TEXT_RE = re.compile(
    rf"(?:{_MONTH}\s+)?(?:19|20)\d{{2}}\s*(?:-|–|—|to)\s*(?:(?:{_MONTH}\s+)?(?:19|20)\d{{2}}|present|current|now)",
    re.I,
)
_JOB_HEADER_RE = re.compile(
    rf"(?:(?<=^)|(?<=[.!?])\s+)(?P<header>[A-Z][^•\n]{{8,190}}?)\s+(?P<dates>{_DATE_RANGE_TEXT_RE.pattern})",
    re.I,
)


def _extract_links_from_text(text: str) -> list[str]:
    raw_links = re.findall(
        r"https?://[^\s)>\]]+|www\.[^\s)>\]]+|(?:linkedin|github)\.com/[^\s)>\]]+|mailto:[^\s)>\]]+",
        text or "",
        flags=re.I,
    )
    out: list[str] = []
    seen: set[str] = set()
    for raw in raw_links:
        link = _clean_link(raw)
        key = link.lower()
        if link and key not in seen:
            seen.add(key)
            out.append(link)
    return out[:20]


def _split_experience_bullets(text: str) -> list[str]:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    normalized = re.sub(r"\s*[•●▪‣◦∙·]\s*", "\n", normalized)
    parts = []
    for chunk in normalized.splitlines():
        chunk = chunk.strip(" -–—•·\t")
        if len(chunk) >= 12:
            parts.append(chunk)
    return parts[:8]


def _experience_entries(lines: list[str]) -> list[dict]:
    """Split resume experience into company/title/date records.

    PDF extraction often merges adjacent roles into one line. We identify role
    headers by their date ranges, then take the text between headers as bullets.
    """
    blob = " ".join(str(line or "").strip() for line in lines if str(line or "").strip())
    blob = re.sub(r"\s+", " ", blob).strip()
    # Some PDF/export paths split years into two tokens ("Jan 20 25"). Repair
    # that before identifying role headers, otherwise later jobs get merged
    # into the previous role's bullet text.
    blob = re.sub(r"\b(19|20)\s+(\d{2})\b", r"\1\2", blob)
    if not blob:
        return []
    matches = list(_JOB_HEADER_RE.finditer(blob))
    entries: list[dict] = []
    for idx, match in enumerate(matches):
        header = re.sub(r"\s+", " ", match.group("header")).strip(" -–—|")
        dates = re.sub(r"\s+", " ", match.group("dates")).strip()
        next_start = matches[idx + 1].start() if idx + 1 < len(matches) else len(blob)
        body = blob[match.end():next_start].strip()
        parts = [p.strip(" -–—") for p in re.split(r"\s*\|\s*|\s*[·•]\s*|\s+[—–]\s+", header) if p.strip(" -–—")]
        title = parts[0] if parts else header
        company = parts[1] if len(parts) > 1 else ""
        location = " | ".join(parts[2:]) if len(parts) > 2 else ""
        if not company and " - " in title:
            title_bits = [p.strip() for p in title.split(" - ", 1)]
            title = title_bits[0]
            company = title_bits[1] if len(title_bits) > 1 else ""
        # Resumes use both "Title — Company" and "Company — Title". If the
        # second segment reads like a job title and the first doesn't, swap so
        # Past Companies shows real employers with real titles.
        if company and _TITLE_WORDS.search(company) and not _TITLE_WORDS.search(title):
            title, company = company, title
        bullets = _split_experience_bullets(body)
        entries.append({
            "title": title[:120],
            "company": company[:120],
            "location": location[:120],
            "dates": dates,
            "duration": dates,
            "duration_label": _duration_label_from_dates(dates),
            "bullets": bullets,
        })
    return entries[:10]


def _experience_lines_from_entries(entries: list[dict], fallback: list[str]) -> list[str]:
    if not entries:
        return fallback[:30]
    lines: list[str] = []
    for entry in entries:
        header = " | ".join(
            part for part in (
                entry.get("title"),
                entry.get("company"),
                entry.get("location"),
                entry.get("duration"),
            )
            if part
        )
        if header:
            lines.append(header)
        for bullet in entry.get("bullets") or []:
            lines.append(f"• {bullet}")
    return lines[:50]


def resume_json_looks_shattered(resume_json: dict) -> bool:
    """Detect stored resume JSON built from word-per-line extractions.

    Older uploads persisted shattered sections (each list item a single
    word). Callers use this to re-derive from parsed_text with the fixed
    pipeline instead of rendering broken documents.
    """
    if not isinstance(resume_json, dict):
        return False
    items: list[str] = []
    for key in ("experience", "education", "projects", "certifications"):
        value = resume_json.get(key)
        if isinstance(value, list):
            items.extend(str(v) for v in value)
    sections = resume_json.get("sections")
    if isinstance(sections, dict):
        for value in sections.values():
            if isinstance(value, list):
                items.extend(str(v) for v in value)
    items = [it for it in items if it.strip()]
    if len(items) < 12:
        return False
    single_word = sum(1 for it in items if len(it.split()) <= 2 and len(it) <= 24)
    return single_word / len(items) >= 0.6


def resume_text_to_json(text: str, *, metadata: Optional[dict] = None) -> dict:
    """Convert parsed resume text into a stable private JSON shape."""
    cleaned = _clean_resume_text(text or "")
    sections = _section_lines(cleaned)
    try:
        from app.utils.text_processing import extract_relevant_keywords, extract_skills_from_text

        skills = extract_skills_from_text(cleaned)
        keywords = extract_relevant_keywords(cleaned, top_n=60)
    except Exception:
        skills = []
        keywords = []

    metadata = metadata or {}
    metadata_links = [str(link).strip() for link in (metadata.get("links") or []) if str(link).strip()]
    urls: list[str] = []
    seen_links: set[str] = set()
    for link in [*metadata_links, *_extract_links_from_text(cleaned)]:
        clean = _clean_link(link)
        key = clean.lower()
        if clean and key not in seen_links:
            seen_links.add(key)
            urls.append(clean)
    phone = _first_match(r"(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}", cleaned)
    email = _first_match(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", cleaned)

    # Full summary paragraph (previously truncated to the first 4 lines,
    # which showed just the candidate's name on shattered resumes).
    summary = " ".join(sections.get("summary", []))[:1200]
    experience_details = _experience_entries(sections.get("experience", []))
    experience_lines = _experience_lines_from_entries(experience_details, sections.get("experience", []))

    return {
        "schema_version": "placeup_resume_v1",
        "contact": {
            "email": email,
            "phone": phone,
            "links": urls[:10],
        },
        "header": sections.get("header", [])[:6],
        "summary": summary,
        "skills": sorted(set(skills)),
        "keywords": keywords,
        "sections": sections,
        "experience": experience_lines,
        "experience_details": experience_details,
        "education": sections.get("education", [])[:20],
        "projects": sections.get("projects", [])[:20],
        "certifications": sections.get("certifications", [])[:20],
        "metadata": metadata,
    }
